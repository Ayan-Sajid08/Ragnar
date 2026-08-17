from io import BytesIO

from docx import Document

from services.chunker import chunk_text


def extract_text(file_bytes: bytes):
    document = Document(BytesIO(file_bytes))

    sections = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # Extract tables
    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            row_text = " | ".join(
                cell for cell in cells if cell
            )

            if row_text:
                sections.append(row_text)

    text = "\n".join(sections)

    if not text.strip():
        return [], False

    return chunk_text(text), False